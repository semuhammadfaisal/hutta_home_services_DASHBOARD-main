from pathlib import Path
from zipfile import ZipFile
import json
from docx import Document
from lxml import etree

src = Path(r"C:\Users\AGC\Downloads\Smplfix.com.docx")
doc = Document(src)

out = {
    "paragraphs": [],
    "tables": [],
    "sections": [],
    "headers": [],
    "footers": [],
    "inline_shapes": len(doc.inline_shapes),
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        out["paragraphs"].append({"index": i, "style": p.style.name, "text": text})

for ti, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    out["tables"].append({"index": ti, "rows": rows})

for i, s in enumerate(doc.sections):
    out["sections"].append({
        "index": i,
        "page_width": s.page_width,
        "page_height": s.page_height,
        "top_margin": s.top_margin,
        "bottom_margin": s.bottom_margin,
        "left_margin": s.left_margin,
        "right_margin": s.right_margin,
    })
    out["headers"].append([p.text for p in s.header.paragraphs if p.text.strip()])
    out["footers"].append([p.text for p in s.footer.paragraphs if p.text.strip()])

ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
with ZipFile(src) as z:
    names = set(z.namelist())
    out["has_comments"] = "word/comments.xml" in names
    out["has_footnotes"] = "word/footnotes.xml" in names
    out["has_endnotes"] = "word/endnotes.xml" in names
    document_xml = etree.fromstring(z.read("word/document.xml"))
    out["tracked_insertions"] = len(document_xml.xpath(".//w:ins", namespaces=ns))
    out["tracked_deletions"] = len(document_xml.xpath(".//w:del", namespaces=ns))
    out["hyperlinks"] = len(document_xml.xpath(".//w:hyperlink", namespaces=ns))
    out["page_breaks"] = len(document_xml.xpath(".//w:br[@w:type='page']", namespaces=ns))
    out["drawings"] = len(document_xml.xpath(".//w:drawing", namespaces=ns))
    if out["has_comments"]:
        comments_xml = etree.fromstring(z.read("word/comments.xml"))
        out["comments"] = [
            "".join(c.xpath(".//w:t/text()", namespaces=ns))
            for c in comments_xml.xpath(".//w:comment", namespaces=ns)
        ]

print(json.dumps(out, ensure_ascii=False, indent=2, default=str))
